# Databricks notebook source
# MAGIC %pip install python-docx

# COMMAND ----------

import os
import pandas as pd
import shutil
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import Cm
from utils.enviroment import BASE_DIR


def get_image_path(atg_code, pos="in", size="xl",
                   root="/Volumes/lc_prd/ml_data_preproc_silver/image/raw_product_images/") -> str:
    """Get image path for given atg"""
    return os.path.join(root, size, atg_code[0], atg_code[1], atg_code[2], f"{atg_code}_{pos}_{size}.jpg")


output = pd.read_csv(os.path.join(BASE_DIR, "output_ranked.csv"))

document = Document()

# narrow margins
sections = document.sections
for section in sections:
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

for cluster in output["cluster"].unique():
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.text = cluster

    paragraph.style = "Heading1"
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run.bold = True
    run.font.size = Pt(16)

    # add table
    table = document.add_table(rows=1, cols=4)
    table.style = "TableGrid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Rank"
    hdr_cells[1].text = "ATG Code"
    hdr_cells[2].text = "Image"
    hdr_cells[3].text = "Details"

    subdf = output[output["cluster"] == cluster].sort_values(by="rank")

    for index, row in subdf.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row["rank"])
        row_cells[1].text = row["atg_code"]
        try:
            row_cells[2].add_paragraph().add_run().add_picture(get_image_path(row["atg_code"]), width=Inches(2.5))
        except Exception as e:
            print(e)
        details = f"""
            Brand: {row["brand_desc"].upper()}\n
            Category: {row["category"]}\n
            Main Accords: {row["main_accords"]}\n
            Sillage: {row["sillage"]}\n
            Longevity: {row["longevity"]}\n
            Tops notes: {row["top_notes"]}\n
            Middle Notes: {row["middle_notes"]}\n
            Base Notes: {row["base_notes"]}\n
        """
        row_cells[3].text = details
    
    cell_widths = [1.5, 2, 5, 8]  # in centimeters
    # Set the width of each cell
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(cell_widths[i])

# direct save to blob is not supported
doc_name = "visualize_rank.docx"
os.makedirs("/tmp/hof_docx/", exist_ok=True)
document.save(f"/tmp/hof_docx/{doc_name}")

dbutils.fs.cp(
    f"file:/tmp/hof_docx/{doc_name}", BASE_DIR.replace("/dbfs", ""), recurse=True
)
shutil.rmtree("/tmp/hof_docx/")


# COMMAND ----------


